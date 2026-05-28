from __future__ import annotations

import io

from fielddesk_worker.parsing.base import ParseError, ParsedSegment


# python-docx exposes heading levels as style names like "Heading 1" .. "Heading 9".
_HEADING_STYLE_PREFIX = "Heading "


def parse_docx(content: bytes) -> list[ParsedSegment]:
    """DOCX parser. Emits one segment per "section" — a paragraph trail between
    Heading-styled paragraphs — preserving the heading path for citations.

    Out of scope for v1: tables-as-structured-data (we extract their text but
    don't preserve cell coordinates), images, embedded objects. A truly empty
    document returns []; a corrupt one raises ParseError.
    """
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as e:
        raise ParseError(f"python-docx is required for DOCX parsing: {e}") from e

    try:
        doc = docx.Document(io.BytesIO(content))
    except Exception as e:  # noqa: BLE001
        raise ParseError(f"could not read docx: {e}") from e

    segments: list[ParsedSegment] = []
    heading_stack: list[str] = []
    current_body: list[str] = []

    def flush(path: list[str]) -> None:
        body = "\n".join(current_body).strip()
        if body:
            segments.append(ParsedSegment(text=body, heading_path=list(path)))
        current_body.clear()

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        style_name = (paragraph.style.name or "") if paragraph.style else ""
        if style_name.startswith(_HEADING_STYLE_PREFIX) and text:
            flush(heading_stack)
            level_str = style_name[len(_HEADING_STYLE_PREFIX):].strip()
            try:
                level = max(1, int(level_str))
            except ValueError:
                level = 1
            del heading_stack[level - 1 :]
            heading_stack.append(text)
        else:
            if text:
                current_body.append(text)
    flush(heading_stack)

    # python-docx skips table cell text in the paragraph iterator; pull it
    # in as additional segments so a tables-only DOCX still produces chunks.
    # We deliberately don't try to reconstruct cell coordinates in v1.
    for table_index, table in enumerate(doc.tables):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text]
            if cells:
                rows.append(" | ".join(cells))
        joined = "\n".join(rows).strip()
        if joined:
            segments.append(
                ParsedSegment(
                    text=joined,
                    heading_path=list(heading_stack),
                    source_locator={"table_index": table_index},
                )
            )
    return segments
